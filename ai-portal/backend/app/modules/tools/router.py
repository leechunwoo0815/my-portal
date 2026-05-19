"""AI工具API路由 - PDF摘要、Markdown转Word、代码解释"""
import os
import re
import uuid
from typing import Any

from fastapi import APIRouter, Depends, status, UploadFile, File
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.deps import get_db, get_current_user
from app.core.exceptions import FileError
from app.modules.tools.schemas import (
    PDFSummaryRequest,
    PDFSummaryResponse,
    MdToWordRequest,
    CodeExplainRequest,
    CodeExplainResponse,
)
from app.services.llm_service import llm_service

router = APIRouter(tags=["AI工具"])


@router.post("/pdf-summary", response_model=PDFSummaryResponse)
def pdf_summary(
    file: UploadFile = File(...),
    max_length: int = 500,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
) -> dict[str, Any]:
    """PDF摘要生成器"""
    if not file.filename.lower().endswith(".pdf"):
        raise FileError("仅支持PDF文件")

    upload_dir = settings.UPLOAD_DIR
    os.makedirs(upload_dir, exist_ok=True)
    safe_filename = f"{uuid.uuid4().hex}_{os.path.basename(file.filename)}"
    safe_filename = re.sub(r'[^\w\-.]', '_', safe_filename)
    file_path = os.path.join(upload_dir, safe_filename)

    with open(file_path, "wb") as f:
        f.write(file.file.read())

    try:
        import fitz
        doc = fitz.open(file_path)
        page_count = len(doc)
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()
        word_count = len(full_text)

        summary_prompt = (
            f"请为以下文档生成一份中文摘要，字数控制在{max_length}字以内。"
            f"摘要应包含文档的核心观点、关键结论和重要数据。\n\n"
            f"文档内容：\n{full_text[:8000]}\n\n"
            f"摘要："
        )

        summary = llm_service.chat_completion(
            messages=[{"role": "user", "content": summary_prompt}],
            model_id=settings.DEFAULT_MODEL,
            temperature=0.5,
            max_tokens=max_length,
            db=db,
        )

        return {
            "summary": summary,
            "page_count": page_count,
            "word_count": word_count,
        }
    except Exception as e:
        raise FileError("PDF处理失败，请稍后重试")
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)


@router.post("/md-to-word")
def md_to_word(
    request: MdToWordRequest,
    current_user = Depends(get_current_user),
) -> dict[str, Any]:
    """Markdown转Word"""
    try:
        from docx import Document
        doc = Document()
        lines = request.markdown.split("\n")
        for line in lines:
            line = line.rstrip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1. ") or line.startswith("2. "):
                doc.add_paragraph(line[3:], style="List Number")
            else:
                doc.add_paragraph(line)

        upload_dir = settings.UPLOAD_DIR
        os.makedirs(upload_dir, exist_ok=True)
        safe_filename = re.sub(r'[^\w\-.]', '_', os.path.basename(request.filename))
        if not safe_filename.endswith('.docx'):
            safe_filename += '.docx'
        output_path = os.path.join(upload_dir, safe_filename)
        doc.save(output_path)

        return {
            "success": True,
            "filename": safe_filename,
            "download_url": f"/uploads/{safe_filename}",
            "message": "转换成功",
        }
    except Exception as e:
        raise FileError("转换失败，请稍后重试")


@router.post("/code-explain", response_model=CodeExplainResponse)
def code_explain(
    request: CodeExplainRequest,
    current_user = Depends(get_current_user),
) -> dict[str, Any]:
    """代码解释器"""
    level_desc = {
        "brief": "简要",
        "detailed": "详细",
        "expert": "专家级",
    }.get(request.detail_level, "详细")
    lang_hint = f"（编程语言：{request.language}）" if request.language else ""

    explain_prompt = (
        f"请对以下代码进行{level_desc}解释{lang_hint}。\n\n"
        f"要求：\n"
        f"1. 说明代码的整体功能和目的\n"
        f"2. 逐段解释关键逻辑\n"
        f"3. 列出关键点（如算法、设计模式等）\n"
        f"4. 提供改进建议（如有）\n\n"
        f"代码：\n```\n{request.code}\n```\n\n"
        f"请按以下JSON格式输出：\n"
        f'{{"explanation": "整体解释", "key_points": ["要点1", "要点2"], "suggestions": ["建议1", "建议2"]}}'
    )

    try:
        import json
        content = llm_service.chat_completion(
            messages=[{"role": "user", "content": explain_prompt}],
            model_id=settings.DEFAULT_MODEL,
            temperature=0.3,
            max_tokens=2048,
        )

        try:
            if "```json" in content:
                json_str = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                json_str = content.split("```")[1].split("```")[0].strip()
            else:
                json_str = content.strip()
            result = json.loads(json_str)
            return {
                "explanation": result.get("explanation", content),
                "key_points": result.get("key_points", []),
                "suggestions": result.get("suggestions", []),
            }
        except json.JSONDecodeError:
            return {
                "explanation": content,
                "key_points": [],
                "suggestions": [],
            }
    except Exception as e:
        raise FileError("代码解释失败，请稍后重试")
